import pdfplumber
import docx
import re
import string


def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file):
    document = docx.Document(file)
    text = "\n".join([para.text for para in document.paragraphs])

    # Also pull text from tables since some resumes use them
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    return text


def extract_text(file, filename):
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def clean_text(text):
    text = text.lower()
    text = re.sub(r"[\n\r\t]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def check_formatting_issues(file, filename, raw_text):
    """Flag common ATS-unfriendly formatting issues."""
    issues = []

    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                images = page.images
                if images:
                    issues.append(
                        "Resume contains images/graphics. Some ATS systems cannot read text inside images (e.g. icons next to contact info)."
                    )
                    break

    if filename.lower().endswith(".docx"):
        document = docx.Document(file)
        if document.tables:
            issues.append(
                "Resume uses tables for layout. Some ATS parsers struggle to read content inside tables in the correct order."
            )

    word_count = len(raw_text.split())
    if word_count < 150:
        issues.append(
            "Resume text content seems very short. Make sure your PDF/DOCX isn't mostly images or icons replacing text."
        )

    if not re.search(r"[\w\.-]+@[\w\.-]+", raw_text):
        issues.append("No email address detected. Make sure your contact info is in plain text, not an image.")

    if not re.search(r"(\+?\d[\d\-\s\(\)]{7,}\d)", raw_text):
        issues.append("No phone number detected in plain text.")

    common_section_headers = ["experience", "education", "skills", "projects", "summary", "objective"]
    found_sections = [h for h in common_section_headers if h in raw_text.lower()]
    if len(found_sections) < 2:
        issues.append(
            "Couldn't detect standard section headers (Experience, Education, Skills, etc.). "
            "Use clear, standard section titles so ATS systems can categorize your content."
        )

    return issues
